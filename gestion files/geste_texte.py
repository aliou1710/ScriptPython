import os;
import os.path;
import glob;
import shutil; 
import re;
import unicodedata;

from docx import Document;

#instantiate document

#document.save('cv.docx') allow to create word file.
#document = Document()
formatsfile =("txt")
source = r'D:\\Cours_bloc_MA2\Langage script python\gestion_textes'
destination= r"D:\\Cours_bloc_MA2\Langage script python\destination"
list_of_element_Todelete = ["-----","Date","Note","Titre"]
print(os.getcwd())

#1 sorted file by date and times
def sortedfilesdirectory():
    list_of_files = os.listdir(source)
    print(list_of_files)  
    #sort by modification time : list_of_files.sort(key=lambda x:os.path.getmtime(x)) 
    #sort by creation time
    list_of_files.sort(key=lambda x:os.path.getctime(x)) 
    last_lists = []
    for item in list_of_files:
        path = os.path.join(source,item)
        last_lists.append(path)
    print(last_lists)
    return last_lists

#2 create folder by name of classes (cours)
def createFolder(listof_folder,destinations):
    print("************creation folder ************")  
    for x in listof_folder:
        path = os.path.join(destinations,x)
        if (os.path.isdir(path)):
            print("exist déjà")
            continue
            
        else:
          
          print(path)
          os.mkdir(path)

#create one flder
def createOnFolder(dirTitle,destinations):
    path = destinations+"\\"+dirTitle
    print(path)
    if(os.path.exists(path)):
        print("exist déjà ")
        
    else:
       os.mkdir(path)

    
def createFile(dirTitle,destinations):

    filepath = destinations+"\\"+dirTitle+"\\"+dirTitle+".docx"
    print(filepath)
    if(os.path.exists(filepath)):
        print("exist déjà ")
        
    else:
       Document().save(filepath)

#delete /n
def deleteBackslashN(lists):
    modifiedList=[]
    for item in lists:
      #eviter la duplication
      if item not in modifiedList:
        if item[-1] == '\n':
            modifiedList.append(item[:-1])
        else:
            modifiedList.append(item)
    return modifiedList


#open file
text = '1.txt'
def method1OpenFile(fileToOpen):
    listes=[]
    with open(fileToOpen,'r',errors='ignore') as file_handler:
        file_handler = open(fileToOpen,'r')
        list_handler = file_handler.readlines()
        listes = deleteBackslashN(list_handler)
        print(listes)
        if listes[-1].lstrip()=="#lu":
            print("fichier lu déjà")
    return listes

#ajouter lu si un ficheir à été lu
def addReadTofile(fileToOpen):
    with open(fileToOpen,'a',errors='ignore') as file_handler:
        raw_data = file_handler.readlines()
        if "#lu" not in raw_data:
            file_handler.write("#lu\n")


def searchTitleToCreateFolder(itemTile):
    lists = []
    with open(itemTile,'r',errors='ignore') as file_handler:
        checkoutlist = file_handler.readlines()
        all_lines =deleteBackslashN(checkoutlist)
      
        for line in all_lines:
            if line.startswith('Titre') :
                print(line)
                phrases= line.split(':')
                lists.append(phrases[1].strip())
                print("phrase : ",phrases[1].strip())
                del checkoutlist
                break
        del all_lines
    return lists    



#search date 
def searchDate(fileToOpenlist):
    msg=""

  
    for item in  fileToOpenlist:
      
        if item.startswith("Date"):
            m = re.split(r':',item)
            msg = m[1].strip()
            break
    return msg



def sendEachLineToWordFile(fileToOpen,Title,Destinations):
#def sendEachLineToWordFile(fileToOpen,deleteElementlist):
   

    with open(fileToOpen,'r',errors='ignore',encoding='ascii') as file_handler:
       
      raw_data = deleteBackslashN(file_handler.readlines())
      date_ = searchDate(raw_data)
       # raw_data=decodeUtf8(raw_data)
        #search error �\xa0
        
      raw_data =[els.replace(u'\xa0','a') for els in raw_data]
        
      print(raw_data)
        
      checkifreadfile="#lu"
      if checkifreadfile not in  raw_data:
        #print(raw_data)
        #la limite des notes ecrites
        indexOfLastElement = raw_data.index('-----')
        print("index of last : ",indexOfLastElement)
        key_Notes = raw_data[3:indexOfLastElement]
        key_Liens = raw_data[indexOfLastElement+1:]
        print(key_Notes)
        print(key_Liens)
       
        filepath = Destinations+"\\"+Title+"\\"+Title+".docx"
        document = Document(filepath)
        
         #set font
        style = document.styles['Normal']
        style.font.name="Calibri"
        for item in  key_Notes:

            # delete line that contains element of list (deleteElementlist)
    
            #if line begin by (**) => it's a heading
            if  item.startswith("**") and item.strip().endswith("**"):
                        
                linefile = item.strip()
                
                res = re.sub(r'\*(.*?)\*', "", linefile)
                print(res)
                document.add_heading(res,0)

            elif item.startswith("-"):
                linefile = item.strip()
               
                res = re.sub(r'\-(.*?)', "", linefile)
                print(res)
                document.add_paragraph(res,style='List Bullet')
            else:
                document.add_paragraph(item)


        for elm in key_Liens:

            if elm.strip().startswith("Liens:"):
                p =document.add_paragraph(elm)
                p.add_run('Liens:').bold =True
            else:
                document.add_paragraph(elm,style='List Bullet')
        document.add_paragraph(date_)

    
        if os.path.exists(filepath):
        #document.save("D:\\Cours_bloc_MA2\Langage script python\destination\Developpement Applications Reseau\Developpement Applications Reseau.docx")
            document.save(filepath)    
           


def conception(formatsfiles_,listfiles,destinations):
   
    print("************open  files ************")  
    for itemFormat  in formatsfiles_:
     
        for  item in listfiles:
             #print("foreach files in sorted files")
             if item.endswith(itemFormat):
                print(item)
                ListTitle=searchTitleToCreateFolder(item.strip())
                print("*****lists*****")
                #loop to create folder and docx file
                Title = ListTitle[0]
                
                #print("title: ",Title,"|  filepath: ",destination+'\\'+Title+"\\"+Title+".docx")
                print("title ", Title," | liens",item)
                print()
                createOnFolder(Title,destinations)
                createFile(Title,destinations)
                sendEachLineToWordFile(item,Title,destinations)
                addReadTofile(item)
                



print() 
listOfFiles_ = sortedfilesdirectory()  
conception(formatsfile,listOfFiles_,destination)       


